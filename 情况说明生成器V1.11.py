import os
import sys
import json
from click import style
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from datetime import datetime
import tkinter as tk
from tkinter import messagebox, ttk, scrolledtext, filedialog, simpledialog, font
import uuid
import pickle
import re
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32print
import win32api
import tempfile
import subprocess
import win32com.client
import threading
import math
import time
import pythoncom  # 添加pythoncom模块导入
import random
import platform

class LoadingDialog(tk.Toplevel):
    """创建一个带有加载动画的对话框"""
    def __init__(self, parent, title="处理中", message="正在处理，请稍候...", size=30):
        super().__init__(parent)
        self.title(title)
        self.resizable(False, False)
        self.configure(bg='white')
        self.transient(parent)  # 设置为父窗口的临时窗口
        self.grab_set()  # 模态窗口
        self.focus_set()  # 获取焦点
        
        # 居中显示
        window_width = 300
        window_height = 150
        position_x = parent.winfo_x() + (parent.winfo_width() - window_width) // 2
        position_y = parent.winfo_y() + (parent.winfo_height() - window_height) // 2
        self.geometry(f"{window_width}x{window_height}+{position_x}+{position_y}")
        
        # 创建加载动画容器
        canvas_frame = tk.Frame(self, bg='white')
        canvas_frame.pack(pady=10)
        
        # 创建画布用于绘制动画
        self.canvas = tk.Canvas(canvas_frame, width=size, height=size, bg='white', highlightthickness=0)
        self.canvas.pack()
        
        # 添加消息标签
        self.message_label = ttk.Label(self, text=message, font=('黑体', 12), background='white')
        self.message_label.pack(pady=10)
        
        # 初始化动画参数
        self.size = size
        self.angle = 0
        self.running = True
        
        # 开始动画
        self.draw_spinner()
    
    def draw_spinner(self):
        """绘制旋转动画"""
        if not self.running:
            return
            
        # 清除画布
        self.canvas.delete("all")
        
        # 计算中心点
        center = self.size / 2
        
        # 计算外圆半径和内圆半径
        outer_radius = center - 5
        inner_radius = outer_radius * 0.6
        
        # 绘制12个点，角度不同透明度不同
        for i in range(12):
            # 计算点的角度（弧度）
            point_angle = math.radians(self.angle + i * 30)
            
            # 计算透明度（最前面的点最不透明）
            alpha = int(255 * (1 - i / 12))
            color = f"#{alpha:02x}{alpha:02x}{alpha:02x}"
            
            # 计算点的坐标
            x = center + outer_radius * math.cos(point_angle)
            y = center + outer_radius * math.sin(point_angle)
            
            # 绘制圆点
            self.canvas.create_oval(x-3, y-3, x+3, y+3, fill=color, outline="")
        
        # 更新角度
        self.angle = (self.angle + 10) % 360
        
        # 继续动画
        self.after(50, self.draw_spinner)
    
    def update_message(self, message):
        """更新提示消息"""
        self.message_label.config(text=message)
    
    def stop(self):
        """停止动画"""
        self.running = False
        self.destroy()

class SituationReportGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("情况说明生成器 V1.11")
        self.root.geometry("1054x1054")  # 设置窗口大小为1054*1054
        
        # 设置程序图标
        try:
            self.root.iconbitmap('./config/icon.ico')
        except:
            pass
            
        # 初始化错误提示记录集合
        self.id_error_shown = set()
        
        # 配置应用程序字体样式
        self.root.option_add('*Font', '黑体 14')
        
        # 配置菜单字体较大的样式
        self.menu_font = font.Font(family="黑体", size=16, weight="bold")
        self.combo_font = font.Font(family="黑体", size=14)
        
        # 获取程序所在目录，用于保存配置和数据文件
        self.app_dir = os.path.dirname(os.path.abspath(__file__))
        # 创建配置保存目录
        self.config_dir = os.path.join(self.app_dir, "CONFIG")
        if not os.path.exists(self.config_dir):
            os.makedirs(self.config_dir)
            
        # 创建文档保存目录（用户的Documents文件夹）
        self.doc_dir = os.path.join(os.path.expanduser("~"), "Documents", "安置房分配协议")
        if not os.path.exists(self.doc_dir):
            os.makedirs(self.doc_dir)
        
        # 存储控件原始尺寸的字典
        self.original_widget_sizes = {}
        
        # 配置菜单
        self.menu_bar = tk.Menu(root)
        root.config(menu=self.menu_bar)
        
        # 添加文件菜单
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0, font=self.combo_font)
        self.menu_bar.add_cascade(label="文件", menu=self.file_menu, font=self.menu_font)
        
        # 添加使用说明菜单
        self.help_guide_menu = tk.Menu(self.menu_bar, tearoff=0, font=self.combo_font)
        self.menu_bar.add_cascade(label="使用说明", menu=self.help_guide_menu, font=self.menu_font)
        self.help_guide_menu.add_command(label="查看使用说明", command=self.show_readme)
        
        # 添加帮助菜单
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0, font=self.combo_font)
        self.menu_bar.add_cascade(label="帮助", menu=self.help_menu, font=self.menu_font)
        self.help_menu.add_command(label="关于", command=self.show_about)
        
        # 创建顶部按钮区域
        self.top_frame = ttk.Frame(root)
        self.top_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 创建左上角按钮容器
        left_button_frame = ttk.Frame(self.top_frame)
        left_button_frame.pack(side=tk.LEFT)
        
        # 添加"打开文档"按钮到左上角
        open_docs_btn = ttk.Button(
            left_button_frame, 
            text="打开文档", 
            command=self.open_documents_folder,
            style="Large.TButton",
            width=10
        )
        open_docs_btn.pack(side=tk.LEFT, padx=5)
        
        # 创建右上角按钮容器
        right_button_frame = ttk.Frame(self.top_frame)
        right_button_frame.pack(side=tk.RIGHT)
        
        # 加载上次保存的界面比例
        self.scale_factor = self.load_ui_scale()  # 加载保存的比例，默认为1.0
        
        # 添加"调整界面"按钮到右上角
        self.scale_var = tk.StringVar()
        # 将加载的比例转换为显示格式
        scale_percent = f"{int(self.scale_factor * 100)}%"
        self.scale_var.set(scale_percent)
        self.scale_options = ["100%", "90%", "80%", "70%"]
        
        # 配置下拉菜单样式以使用较大字体
        style = ttk.Style()
        style.configure('TMenubutton', font=('黑体', 14))
        
        scale_menu = ttk.OptionMenu(
            right_button_frame,
            self.scale_var,
            self.scale_var.get(),
            *self.scale_options,
            command=self.change_scale
        )
        scale_menu.configure(width=5)
        
        # 设置下拉菜单的字体
        root.option_add('*TCombobox*Listbox.font', self.combo_font)
        
        scale_menu.pack(side=tk.RIGHT, padx=5)
        
        # 创建标签显示"界面比例:"
        scale_label = ttk.Label(right_button_frame, text="界面比例:", font=('黑体', 14))
        scale_label.pack(side=tk.RIGHT, padx=5)
        
        # 创建数据输入区域（占满整个窗口）
        self.data_frame = ttk.Frame(root)
        self.data_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 设置数据输入区域
        self.setup_data_area()
        
        # 创建苹果风格按钮
        self.create_apple_style_button()
        
        # 创建按钮区域 - 移至底部中间
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill=tk.X, pady=25)  # 增加垂直边距
        
        # 创建一个居中的容器
        center_frame = ttk.Frame(button_frame)
        center_frame.pack(expand=True)
        
        # 设置"生成情况说明"按钮 - 变大、苹果风格、居中放置
        generate_btn = ttk.Button(
            center_frame, 
            text="生成情况说明", 
            command=self.generate_report, 
            takefocus=0,
            style="Apple.TButton",
            width=30  # 增大按钮宽度
        )
        generate_btn.pack(side=tk.LEFT, padx=10, pady=10)  # 修改为左对齐，添加水平边距
        
        # 设置"直接打印"按钮 - 与生成按钮样式一致
        print_btn = ttk.Button(
            center_frame, 
            text="直接打印", 
            command=self.print_report, 
            takefocus=0,
            style="Apple.TButton",
            width=30  # 与生成按钮宽度一致
        )
        print_btn.pack(side=tk.LEFT, padx=10, pady=10)  # 左对齐，与生成按钮并排
        
        # 创建状态栏
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 初始化变量
        self.apartment_frames = []
        
        # 数据列表，用于存储可选择的人员信息
        self.person_data = []
        
        # 添加第一个默认公寓表单
        self.add_apartment_frame()
        # 添加第二个默认公寓表单
        self.add_apartment_frame()
        
        # 加载最近文件列表
        self.recent_files = []
        self.load_recent_files()
        self.update_recent_files_menu()
        
        # 设置窗口大小为根据比例计算的值，不允许调整
        base_width = 1054
        base_height = 1054
        new_width = int(base_width * self.scale_factor)
        new_height = int(base_height * self.scale_factor)
        
        # 计算窗口居中显示的位置
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x_position = (screen_width - new_width) // 2
        y_position = (screen_height - new_height) // 2
        
        # 设置窗口大小和位置，格式为"宽x高+x+y"
        self.root.geometry(f"{new_width}x{new_height}+{x_position}+{y_position}")
        self.root.resizable(False, False)  # 禁止调整窗口大小
        
        # 如果加载的比例不是100%，立即应用缩放效果
        if self.scale_factor != 1.0:
            self.apply_scale(self.scale_factor)
            
        # 注册窗口关闭事件，保存当前界面比例
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        
    def load_ui_scale(self):
        """加载保存的界面比例设置"""
        ui_config_path = os.path.join(self.config_dir, "ui_settings.json")
        if os.path.exists(ui_config_path):
            try:
                with open(ui_config_path, 'r', encoding='utf-8') as f:
                    ui_settings = json.load(f)
                    return ui_settings.get("scale_factor", 1.0)
            except:
                return 1.0  # 默认100%比例
        return 1.0
        
    def save_ui_scale(self):
        """保存当前界面比例设置"""
        ui_config_path = os.path.join(self.config_dir, "ui_settings.json")
        try:
            ui_settings = {"scale_factor": self.scale_factor}
            with open(ui_config_path, 'w', encoding='utf-8') as f:
                json.dump(ui_settings, f, ensure_ascii=False)
        except:
            print("保存界面比例设置失败")
            
    def on_close(self):
        """窗口关闭时的处理"""
        # 保存当前界面比例
        self.save_ui_scale()
        # 保存其他配置
        self.save_recent_files()
        # 关闭窗口
        self.root.destroy()
        
    def change_scale(self, value):
        """根据选择的比例调整界面大小和字体"""
        # 解析比例值（去掉百分号并转换为浮点数）
        scale = float(value.replace('%', '')) / 100.0
        
        # 保存当前比例
        old_scale = self.scale_factor
        self.scale_factor = scale
        
        # 应用缩放
        self.apply_scale(scale)
        
        # 更新状态栏提示
        self.status_var.set(f"界面比例已调整为 {value}")
        
    def apply_scale(self, scale):
        """应用缩放比例到所有UI元素"""
        # 计算新的窗口尺寸
        base_width = 1054
        base_height = 1054 
        new_width = int(base_width * scale)
        new_height = int(base_height * scale)
        
        # 获取当前窗口位置
        current_geometry = self.root.geometry()
        # 从geometry字符串中提取位置信息 "widthxheight+x+y"
        position_part = current_geometry.split('+', 1)
        if len(position_part) > 1:
            position = '+' + position_part[1]  # 保留 "+x+y" 部分
        else:
            # 如果没有位置信息，使用当前位置
            position = ''
        
        # 调整窗口大小但保持位置不变
        self.root.geometry(f"{new_width}x{new_height}{position}")
        
        # 先允许窗口调整大小，进行缩放后再禁止调整
        self.root.resizable(True, True)
        
        # 创建或重新配置样式
        style = ttk.Style()
        style.configure(
            "Large.TButton",
            font=("黑体", int(14 * scale)),  # 调整字体
            padding=(int(15 * scale), int(8 * scale)),  # 调整内边距
            width=int(10 * scale)  # 调整宽度，确保文字完整显示
        )
        
        # 根据比例调整字体大小和控件样式
        self.update_font_sizes(scale)
        
        # 调整控件尺寸
        self.update_widget_sizes(scale)
        
        # 应用更改后再次禁止调整窗口大小
        self.root.resizable(False, False)
        
        # 触发一次动态控件的更新 - 确保动态添加的控件也得到正确缩放
        self.update_dynamic_elements(scale)
        
        # 确保"添加人员"按钮文字显示完整
        if hasattr(self, 'add_person_btn'):
            self.add_person_btn.configure(text="添加人员", width=int(10 * scale))
    
    def update_dynamic_elements(self, scale):
        """更新动态添加的控件的缩放比例"""
        # 确保"添加人员"按钮缩放正确
        if hasattr(self, 'add_person_btn'):
            self._update_widget_fonts_recursive(self.add_person_btn, scale)
            self._update_widget_sizes_recursive(self.add_person_btn, scale)
        
        # 更新所有家庭成员行的控件 
        for frame in self.member_frames:
            for child in frame.winfo_children():
                self._update_widget_fonts_recursive(child, scale)
                self._update_widget_sizes_recursive(child, scale)
                
                # 特别处理删除按钮，确保样式一致
                if isinstance(child, ttk.Button) and child['text'] == "删除":
                    child.configure(style='Large.TButton')
        
        # 更新所有公寓行的控件
        for apt in self.apartment_frames:
            # 更新主要控件
            for child in apt["frame"].winfo_children():
                self._update_widget_fonts_recursive(child, scale)
                self._update_widget_sizes_recursive(child, scale)
            
            # 更新所有者容器
            for child in apt["owners_container"].winfo_children():
                self._update_widget_fonts_recursive(child, scale)
                self._update_widget_sizes_recursive(child, scale)
                
                # 特别处理加号和减号按钮，确保样式一致
                if isinstance(child, ttk.Button) and (child['text'] == "+" or child['text'] == "-"):
                    child.configure(style='AddOwner.TButton')
                    
    def update_font_sizes(self, scale):
        """根据缩放比例更新所有控件的字体大小"""
        # 获取所有样式
        style = ttk.Style()
        
        # 定义基础字体大小
        base_font_sizes = {
            'TLabel': 14,
            'Tall.TEntry': 14,
            'Large.TButton': 14,
            'Apple.TButton': 14,
            'TableHeader.TLabel': 14,
            'Table.TEntry': 14,
            'Tall.TCombobox': 14,
            'AddOwner.TButton': 14,
            'TLabelframe.Label': 14,  # 添加LabelFrame标题
            'title': 24,  # 主标题
            'section': 16  # 章节标题
        }
        
        # 更新各样式字体大小
        for style_name, base_size in base_font_sizes.items():
            font_size = int(base_size * scale)
            try:
                style.configure(style_name, font=('黑体', font_size))
            except tk.TclError:
                # 如果样式不存在，跳过
                continue
        
        # 更新下拉列表字体
        self.root.option_add('*TCombobox*Listbox.font', ('黑体', int(14 * scale)))
        
        # 处理特殊控件的字体
        for widget in self.root.winfo_children():
            self._update_widget_fonts_recursive(widget, scale)
        
        # 直接处理主标题
        self._find_and_update_title(self.root, scale)
    
    def _find_and_update_title(self, parent, scale):
        """查找并更新主标题和章节标题的字体"""
        style = ttk.Style()
        
        for widget in parent.winfo_children():
            try:
                # 检查是不是标题文本
                if isinstance(widget, (ttk.Label, tk.Label)):
                    text = widget.cget('text')
                    if text == "集体土地安置房分配协议":
                        # 这是主标题
                        widget.configure(font=("黑体", int(24 * scale), "bold"))
                    elif text in ["家庭成员", "项目信息", "安置达成以下分配协议"]:
                        # 这些是章节标题
                        if "(" in text:  # 如果包含括号（例如"家庭成员(0人)"），保留该部分
                            widget.configure(font=("黑体", int(16 * scale), "bold"))
                        else:
                            widget.configure(font=("黑体", int(16 * scale), "bold"))
                    # 处理表单字段标签 - 更全面的匹配
                    elif ":" in text or text in ["姓名", "身份证号码", "拆迁项目名称", "安置小区名称",
                                               "楼号", "单元号", "房号", "户型面积", "归属人", "操作"]:
                        widget.configure(font=("黑体", int(14 * scale)))
                
                # 处理LabelFrame的标题
                elif isinstance(widget, ttk.LabelFrame):
                    # 创建自定义样式
                    custom_style = f"Custom.TLabelframe.{int(scale*100)}"
                    style.configure(f"{custom_style}.Label", font=('黑体', int(14 * scale)))
                    widget.configure(style=custom_style)
            except (tk.TclError, AttributeError) as e:
                # print(f"Error updating title: {e}")
                pass
            
            # 递归处理子控件
            if hasattr(widget, 'winfo_children'):
                self._find_and_update_title(widget, scale)
    
    def _update_widget_fonts_recursive(self, widget, scale):
        """递归更新所有控件的字体大小"""
        try:
            # 检查是否为标签类型控件并且有特殊字体
            if isinstance(widget, (ttk.Label, tk.Label)):
                font = widget.cget('font')
                if font:
                    if isinstance(font, tuple):
                        family, size = font[0], font[1]
                        new_size = int(size * scale)
                        # 检查是否有额外属性如bold
                        if len(font) > 2:
                            widget.configure(font=(family, new_size) + font[2:])
                        else:
                            widget.configure(font=(family, new_size))
                    elif isinstance(font, str):
                        # 处理字符串形式的字体
                        try:
                            if font.startswith('{'):
                                # 解析复杂字体格式，如 "{黑体} 24 bold"
                                parts = font.strip('{}').split()
                                if len(parts) >= 2:
                                    family = parts[0]
                                    size = float(parts[1])
                                    new_size = int(size * scale)
                                    
                                    # 重建字体字符串
                                    if len(parts) > 2:  # 包含样式，如bold
                                        new_font = f"{{{family}}} {new_size} {' '.join(parts[2:])}"
                                    else:
                                        new_font = f"{{{family}}} {new_size}"
                                    
                                    widget.configure(font=new_font)
                        except:
                            pass
                # 对于没有显式设置字体的标签，应用默认字体缩放
                else:
                    # 检查是否为表单字段标签
                    text = widget.cget('text')
                    if text in ["户主:", "姓名:", "身份证号码:", "拆迁项目名称:", "安置小区名称:", 
                               "楼号", "单元号", "房号", "户型面积", "归属人", "操作"]:
                        widget.configure(font=('黑体', int(14 * scale)))
            
            # 处理Text控件
            elif isinstance(widget, tk.Text):
                font = widget.cget('font')
                if font:
                    if isinstance(font, tuple):
                        family, size = font[0], font[1]
                        new_size = int(size * scale)
                        widget.configure(font=(family, new_size))
            
            # 处理普通tk控件
            elif isinstance(widget, (tk.Button, tk.Entry)):
                font = widget.cget('font')
                if font:
                    if isinstance(font, tuple):
                        family, size = font[0], font[1]
                        new_size = int(size * scale)
                        widget.configure(font=(family, new_size))
        except (tk.TclError, AttributeError):
            # 忽略不支持font属性的控件
            pass
        
        # 递归处理所有子控件
        if hasattr(widget, 'winfo_children'):
            for child in widget.winfo_children():
                self._update_widget_fonts_recursive(child, scale)
    
    def update_widget_sizes(self, scale):
        """更新所有控件的尺寸"""
        # 更新所有ttk控件的尺寸
        style = ttk.Style()
        
        # 更新按钮和输入框的内边距
        padding_configs = {
            'Apple.TButton': (int(20 * scale), int(10 * scale)),
            'Large.TButton': (int(15 * scale), int(8 * scale)),
            'Tall.TEntry': (int(5 * scale), int(10 * scale)),
            'Tall.TCombobox': (int(5 * scale), int(10 * scale)),
            'AddOwner.TButton': (int(0), int(10 * scale))
        }
        
        for style_name, padding in padding_configs.items():
            try:
                style.configure(style_name, padding=padding)
            except tk.TclError:
                continue
        
        # 递归更新所有控件的尺寸
        self._update_widget_sizes_recursive(self.root, scale)
    
    def _update_widget_sizes_recursive(self, widget, scale):
        """递归更新所有控件尺寸"""
        # 跳过+/-号按钮的尺寸调整，保持固定大小
        if isinstance(widget, ttk.Button) and widget['text'] in ["+", "-"]:
            return
            
        widget_id = str(widget)
        
        # 根据控件类型调整尺寸
        if isinstance(widget, ttk.Entry) or isinstance(widget, ttk.Combobox):
            # 检查是否有原始尺寸信息
            if hasattr(self, 'original_widget_sizes') and widget_id in self.original_widget_sizes:
                original_width = self.original_widget_sizes[widget_id]
                widget.configure(width=int(original_width * scale))
                
        # 特殊类型的标签和按钮也需要调整尺寸
        elif isinstance(widget, ttk.Label) and hasattr(widget, 'configure') and 'width' in widget.configure():
            if hasattr(self, 'original_widget_sizes') and widget_id in self.original_widget_sizes:
                original_width = self.original_widget_sizes[widget_id]
                widget.configure(width=int(original_width * scale))
                
        elif isinstance(widget, ttk.Button) and widget['text'] == "删除":
            # 删除按钮使用固定尺寸
            if hasattr(self, 'original_widget_sizes') and widget_id in self.original_widget_sizes:
                original_width = self.original_widget_sizes[widget_id]
                widget.configure(width=int(original_width * scale))
        
        # 递归处理子控件
        for child in widget.winfo_children():
            self._update_widget_sizes_recursive(child, scale)
    
    def create_apple_style_button(self):
        """创建苹果风格的按钮样式"""
        style = ttk.Style()
        
        # 配置苹果风格按钮
        style.configure(
            'Apple.TButton',
            font=('黑体', 14),  # 增大字体
            padding=(20, 10),  # 增大内边距
            relief='flat',
            background='#0078d7',
            foreground='white'
        )
        
        # 设置鼠标悬停效果
        style.map(
            'Apple.TButton',
            background=[('active', '#005bb5'), ('disabled', '#cccccc')],
            foreground=[('disabled', '#999999')]
        )
        
        # 配置大号按钮样式 - 增加固定宽度
        style.configure(
            "Large.TButton",
            font=("黑体", 14),  # 修改字体
            padding=(15, 8),  # 增大内边距
            width=10  # 增加宽度，确保文字完整显示
        )
        
        # 配置输入框样式，保证所有输入框高度一致
        style.configure(
            "Tall.TEntry",
            padding=(5, 10),  # 增大内边距，确保与表格中的输入框一致
            font=("黑体", 14)  # 统一字体
        )
        
        # 配置下拉框中的列表样式（设置字体大小）
        self.root.option_add('*TCombobox*Listbox.font', ('黑体', 14))  # 增大下拉列表字体
        
        # 为了确保按钮样式正确显示，我们需要覆盖一些默认的ttk样式
        style.layout("Apple.TButton", [
            ("Button.padding", {
                "children": [
                    ("Button.label", {"sticky": "nswe"})
                ],
                "sticky": "nswe"
            })
        ])
        
        # 设置按钮内部填充
        style.configure("Apple.TButton.padding", padding=8)
    
    def load_recent_files(self):
        """加载最近文件列表"""
        recent_files_path = os.path.join(self.config_dir, "recent_files.json")
        if os.path.exists(recent_files_path):
            try:
                with open(recent_files_path, 'r', encoding='utf-8') as f:
                    self.recent_files = json.load(f)
            except:
                self.recent_files = []
    
    def save_recent_files(self):
        """保存最近文件列表"""
        recent_files_path = os.path.join(self.config_dir, "recent_files.json")
        try:
            with open(recent_files_path, 'w', encoding='utf-8') as f:
                json.dump(self.recent_files, f, ensure_ascii=False)
        except:
            print("保存最近文件列表失败")
    
    def update_recent_files_menu(self):
        """更新最近文件菜单 - 直接显示在文件菜单中"""
        # 完全清空菜单
        self.file_menu.delete(0, 'end')
        
        # 先添加最近文件
        if not self.recent_files:
            self.file_menu.add_command(label="无最近文件", state=tk.DISABLED)
        else:
            # 添加最近文件列表
            for file_path in self.recent_files:
                # 提取文件名，不显示完整路径
                file_name = os.path.basename(file_path)
                self.file_menu.add_command(
                    label=file_name,
                    command=lambda path=file_path: self.load_data_from_file(path)
                )
        
        # 最后添加退出选项
        self.file_menu.add_separator()
        self.file_menu.add_command(label="退出", command=self.root.quit)
    
    def add_to_recent_files(self, file_path):
        """添加文件到最近文件列表"""
        # 如果文件已在列表中，先移除
        if file_path in self.recent_files:
            self.recent_files.remove(file_path)
        
        # 添加到列表最前面
        self.recent_files.insert(0, file_path)
        
        # 限制列表长度为10
        if len(self.recent_files) > 10:
            self.recent_files = self.recent_files[:10]
        
        # 保存列表并更新菜单
        self.save_recent_files()
        self.update_recent_files_menu()
    
    def load_data_from_file(self, file_path):
        """从文件加载数据"""
        try:
            if not os.path.exists(file_path):
                messagebox.showerror("错误", f"文件不存在: {file_path}")
                # 从最近文件列表中移除
                if file_path in self.recent_files:
                    self.recent_files.remove(file_path)
                    self.save_recent_files()
                    self.update_recent_files_menu()
                return
            
            with open(file_path, 'rb') as f:
                data = pickle.load(f)
            
            # 先清空当前表单
            self.clear_form(confirm=False)
            
            # 加载家庭成员
            for member in data.get("members", []):
                # 如果现有行数不够，添加新行
                if len(self.member_name_vars) <= len(data["members"]):
                    self.add_family_member_row()
                
                idx = data["members"].index(member)
                if idx < len(self.member_name_vars):
                    self.member_name_vars[idx].set(member.get("name", ""))
                    self.member_id_vars[idx].set(member.get("id_card", ""))
            
            # 加载项目信息
            self.project_var.set(data.get("project", ""))
            self.community_var.set(data.get("community", ""))
            
            # 加载房产信息
            apartments = data.get("apartments", [])
            # 确保有足够的房产行
            while len(self.apartment_frames) < len(apartments):
                self.add_apartment_frame()
            
            # 填充房产信息
            for i, apt in enumerate(apartments):
                if i < len(self.apartment_frames):
                    apt_frame = self.apartment_frames[i]
                    apt_frame["building"].set(apt.get("building", ""))
                    apt_frame["unit"].set(apt.get("unit", ""))
                    apt_frame["room"].set(apt.get("room", ""))
                    apt_frame["area"].set(apt.get("area", ""))
                    
                    # 设置所有者
                    owners = apt.get("owners", [])
                    # 确保有足够的所有者行
                    while len(apt_frame["owner_vars"]) < len(owners):
                        self.add_owner_to_apartment(apt_frame)
                    
                    # 填充所有者信息
                    for j, owner in enumerate(owners):
                        if j < len(apt_frame["owner_vars"]):
                            apt_frame["owner_vars"][j].set(owner.get("name", ""))
                            apt_frame["id_card_vars"][j].set(owner.get("id_card", ""))
            
            # 更新统计数据
            self.update_person_data()
            
            # 不显示弹窗提示，只更新状态栏
            self.status_var.set(f"已加载数据: {os.path.basename(file_path)}")
            
            # 将此文件添加到最近文件列表的最前面
            self.add_to_recent_files(file_path)
            
        except Exception as e:
            messagebox.showerror("错误", f"加载数据失败: {str(e)}")
            self.status_var.set("加载数据失败")

    def save_data_to_file(self, file_path):
        """保存数据到文件"""
        try:
            # 确保目录存在
            dir_path = os.path.dirname(file_path)
            if not os.path.exists(dir_path):
                os.makedirs(dir_path)
            
            # 收集数据
            data = self.collect_data()
            
            # 序列化数据并保存
            with open(file_path, 'wb') as f:
                pickle.dump(data, f)
            
            # 添加到最近文件列表
            self.add_to_recent_files(file_path)
            
            return True
        except Exception as e:
            messagebox.showerror("错误", f"保存数据失败: {str(e)}")
            return False
    
    def setup_data_area(self):
        """设置数据输入区域，完全匹配Excel模板布局"""
        # 创建主容器框架，使用Canvas和滚动条
        main_container = ttk.Frame(self.data_frame)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # 创建Canvas和垂直滚动条
        self.canvas = tk.Canvas(main_container, bd=0, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # 布局Canvas和滚动条
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 在Canvas上创建一个框架放置内容
        self.data_scroll_frame = ttk.Frame(self.canvas)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.data_scroll_frame, anchor="nw")
        
        # 绑定事件以调整滚动区域
        self.data_scroll_frame.bind("<Configure>", self._configure_scroll_region)
        self.canvas.bind("<Configure>", self._configure_canvas)
        
        # 绑定鼠标滚轮事件
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
        # 主标题
        title_frame = ttk.Frame(self.data_scroll_frame)
        title_frame.pack(fill=tk.X, pady=10)
        
        title_label = ttk.Label(title_frame, text="集体土地安置房分配协议", font=("黑体", 24, "bold"))
        title_label.pack()
        
        # 家庭成员区域 - 包含添加成员按钮
        family_header_frame = ttk.Frame(self.data_scroll_frame)
        family_header_frame.pack(fill=tk.X, padx=15, pady=15)  # 增加内边距
        
        family_title_label = ttk.Label(family_header_frame, text="家庭成员", font=("黑体", 16, "bold"))
        family_title_label.pack(side=tk.LEFT, padx=5)  # 修改字体
        
        # 不再显示人数
        # self.members_count_label = ttk.Label(family_header_frame, text="(0人)", font=("黑体", 16))
        # self.members_count_label.pack(side=tk.LEFT, padx=(0, 5))
        
        # 将"添加人员"按钮移动到家庭成员标签右侧，设置takefocus=0使Tab键跳过此按钮
        self.add_person_btn = ttk.Button(family_header_frame, text="添加人员", command=self.add_family_member, takefocus=0, style='Large.TButton', width=10)
        self.add_person_btn.pack(side=tk.LEFT, padx=5)
        
        # 创建家庭成员容器
        self.family_container = ttk.Frame(self.data_scroll_frame)
        self.family_container.pack(fill=tk.X, padx=15, pady=5)  # 增加内边距
        
        # 家庭成员列表和变量
        self.member_name_vars = []
        self.member_id_vars = []
        self.member_entries = []
        self.member_frames = []
        
        # 默认添加4个家庭成员输入框
        for _ in range(4):
            self.add_family_member_row()
        
        # 项目信息 - 减少内边距
        info_frame = ttk.LabelFrame(self.data_scroll_frame, text="项目信息", padding=8)  # 减少内边距
        info_frame.pack(fill=tk.X, padx=15, pady=10)  # 减少外边距
        
        # 第一行
        ttk.Label(info_frame, text="拆迁项目名称:", font=("黑体", 14)).grid(row=0, column=0, sticky=tk.W, padx=8, pady=5)  # 减少上下间距
        self.project_var = tk.StringVar()
        project_entry = ttk.Entry(info_frame, textvariable=self.project_var, width=45, font=("黑体", 14), style='Tall.TEntry')
        project_entry.grid(row=0, column=1, sticky=tk.W, padx=8, pady=5)  # 减少上下间距
        
        # 第二行 - 修改为"安置小区名称"
        ttk.Label(info_frame, text="安置小区名称:", font=("黑体", 14)).grid(row=1, column=0, sticky=tk.W, padx=8, pady=5)  # 减少上下间距
        self.community_var = tk.StringVar()
        community_entry = ttk.Entry(info_frame, textvariable=self.community_var, width=45, font=("黑体", 14), style='Tall.TEntry')
        community_entry.grid(row=1, column=1, sticky=tk.W, padx=8, pady=5)  # 减少上下间距
        
        # 保留标签但设为空
        self.apartments_count_label = ttk.Label(info_frame, text="", font=("黑体", 14))
        self.apartments_count_label.grid(row=1, column=2, sticky=tk.W, padx=8, pady=5)
        
        # 房产信息区域
        self.property_frame = ttk.LabelFrame(self.data_scroll_frame, text="安置达成以下分配协议", padding=10)
        self.property_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # 表格主容器
        self.table_container = ttk.Frame(self.property_frame)
        self.table_container.pack(fill=tk.BOTH, expand=True)
        
        # 创建表格风格的布局 - 更符合图片风格
        self.create_apartment_table_header()
        
        # 添加房产按钮 - 居中放置，设置takefocus=0使Tab键跳过此按钮
        btn_frame = ttk.Frame(self.property_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        
        add_apt_btn = ttk.Button(btn_frame, text="添加房源", command=self.add_apartment_frame, width=20, takefocus=0, style='Large.TButton')
        add_apt_btn.pack(side=tk.TOP)
    
    def create_apartment_table_header(self):
        """创建房产表格的表头 - 更符合图片样式"""
        # 设置表格样式
        style = ttk.Style()
        # 使用与程序其他部分一致的字体设置
        style.configure("Table.TFrame", background="#ffffff")
        style.configure("TableHeader.TLabel", font=("黑体", 14), anchor="center", background="#f0f0f0")  # 表头字体
        style.configure("TableCell.TEntry", padding=3, font=("黑体", 14))  # 单元格字体
        
        # 设置+/-按钮的样式，不随缩放变化
        style.configure("AddOwner.TButton", font=("黑体", 14, "bold"), padding=(0, 8))  # 按钮字体和内边距
        
        # 设置下拉框样式，增加高度
        style.configure('Tall.TCombobox', padding=(5, 10), font=('黑体', 14))  # 增加下拉框高度与输入框一致
        
        # 表头列名和宽度
        headers = ["楼号", "单元号", "房号", "户型面积", "归属人", "", "身份证号码", "操作"]
        # 修改宽度配置，为"+"按钮添加一列，增大列宽
        widths = [8, 8, 8, 14, 14, 4, 28, 8]  # 增大宽度
        
        # 创建表头行 - 不包含在表格内
        header_frame = ttk.Frame(self.table_container)
        header_frame.pack(fill=tk.X, pady=(0, 0))
        
        # 添加表头标签 - 无边框
        for col, (text, width) in enumerate(zip(headers, widths)):
            # 跳过"+"按钮列的表头及"操作"列表头
            if text == "" or text == "操作":
                continue
                
            # 对于"身份证号码"标题特殊处理
            if text == "身份证号码":
                # 向右移动150像素，增加宽度并使用左对齐，防止文本被截断
                header_label = ttk.Label(header_frame, text=text, width=width+4, style="TableHeader.TLabel", anchor="w")
                header_label.grid(row=0, column=col-1, columnspan=2, sticky="w", padx=(150, 0), pady=0)
            else:
                header_label = ttk.Label(header_frame, text=text, width=width, style="TableHeader.TLabel")
                header_label.grid(row=0, column=col, sticky="nsew", padx=0, pady=0)
            header_frame.columnconfigure(col, weight=0)
        
        # 表格容器框架 - 只包含数据行，不包含标题
        self.table_frame = ttk.Frame(self.table_container, style="Table.TFrame")
        self.table_frame.pack(fill=tk.X, pady=(0, 5))
        
        # 数据行容器
        self.apartments_container = ttk.Frame(self.table_frame, style="Table.TFrame")
        self.apartments_container.pack(fill=tk.X)
        
        # 存储列宽信息供添加行时使用
        self.column_widths = widths
    
    def _on_mousewheel(self, event):
        """处理鼠标滚轮事件"""
        self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    
    def add_family_member_row(self):
        """添加一行家庭成员输入"""
        row_idx = len(self.member_name_vars)
        
        # 创建一个frame来包含这行的输入
        member_frame = ttk.Frame(self.family_container)
        member_frame.pack(fill=tk.X, pady=3)  # 增加垂直间距
        
        name_var = tk.StringVar()
        id_var = tk.StringVar()
        
        # 为不同缩放比例定义精确的控件尺寸
        sizes = {
            1.0: {  # 100%
                'name_label_width': 8,
                'name_entry_width': 18,
                'id_label_width': 14,
                'id_entry_width': 28,
                'delete_btn_width': 10,
                'font_size': 14
            },
            0.9: {  # 90%
                'name_label_width': 7,
                'name_entry_width': 16,
                'id_label_width': 13,
                'id_entry_width': 25,
                'delete_btn_width': 9,
                'font_size': 13
            },
            0.8: {  # 80%
                'name_label_width': 6,
                'name_entry_width': 14,
                'id_label_width': 11,
                'id_entry_width': 22,
                'delete_btn_width': 8,
                'font_size': 11
            },
            0.7: {  # 70%
                'name_label_width': 6,
                'name_entry_width': 13,
                'id_label_width': 10,
                'id_entry_width': 20,
                'delete_btn_width': 7,
                'font_size': 10
            }
        }
        
        # 获取当前缩放比例
        current_scale = getattr(self, 'scale_factor', 1.0)
        
        # 确定当前缩放比例的控件尺寸
        if abs(current_scale - 0.9) < 0.01:  # 90%
            size_key = 0.9
        elif abs(current_scale - 0.8) < 0.01:  # 80%
            size_key = 0.8
        elif abs(current_scale - 0.7) < 0.01:  # 70%
            size_key = 0.7
        else:  # 默认100%
            size_key = 1.0
        
        # 获取当前缩放比例下的尺寸
        dims = sizes[size_key]
        
        # 测量现有输入框的尺寸用于参考（如果有的话）
        if self.member_entries and len(self.member_entries) > 0:
            try:
                # 获取第一个输入框的实际尺寸
                existing_name_entry, existing_id_entry = self.member_entries[0]
                existing_name_width = existing_name_entry.cget('width')
                existing_id_width = existing_id_entry.cget('width')
                
                # 如果存在有效的尺寸，使用它们替代预设值
                if existing_name_width and existing_name_width > 0:
                    dims['name_entry_width'] = existing_name_width
                if existing_id_width and existing_id_width > 0:
                    dims['id_entry_width'] = existing_id_width
            except (tk.TclError, AttributeError, IndexError):
                # 忽略尺寸获取失败的情况
                pass
        
        # 创建控件并应用尺寸
        # 根据行索引决定显示"户主:"还是"姓名:"
        label_text = "户主:" if row_idx == 0 else "姓名:"
        name_label = ttk.Label(member_frame, text=label_text, width=dims['name_label_width'], 
                               font=("黑体", dims['font_size']))
        name_label.pack(side=tk.LEFT, padx=8)
        
        # 确保使用Tall.TEntry样式，统一输入框高度
        style = ttk.Style()
        style.configure('Tall.TEntry', padding=(5, 10), font=('黑体', dims['font_size']))
        
        name_entry = ttk.Entry(member_frame, textvariable=name_var, width=dims['name_entry_width'], 
                               font=("黑体", dims['font_size']), style='Tall.TEntry')
        name_entry.pack(side=tk.LEFT, padx=8)
        
        id_label = ttk.Label(member_frame, text="身份证号码:", width=dims['id_label_width'], 
                             font=("黑体", dims['font_size']), anchor="w")
        id_label.pack(side=tk.LEFT, padx=8)
        
        id_entry = ttk.Entry(member_frame, textvariable=id_var, width=dims['id_entry_width'], 
                             font=("黑体", dims['font_size']), style='Tall.TEntry')
        id_entry.pack(side=tk.LEFT, padx=8)
        
        delete_btn = ttk.Button(member_frame, text="删除", takefocus=0, style='Large.TButton', 
                               width=dims['delete_btn_width'],
                               command=lambda frame=member_frame, idx=row_idx: 
                                        self.delete_family_member(frame, idx))
        delete_btn.pack(side=tk.LEFT, padx=8)
        
        # 存储基准尺寸(100%)用于可能的缩放
        base_dims = sizes[1.0]
        if hasattr(self, 'original_widget_sizes'):
            self.original_widget_sizes[str(name_label)] = base_dims['name_label_width']
            self.original_widget_sizes[str(name_entry)] = base_dims['name_entry_width']
            self.original_widget_sizes[str(id_label)] = base_dims['id_label_width']
            self.original_widget_sizes[str(id_entry)] = base_dims['id_entry_width']
            self.original_widget_sizes[str(delete_btn)] = base_dims['delete_btn_width']
        
        # 存储输入控件
        self.member_entries.append((name_entry, id_entry))
        
        # 当信息改变时，更新人员列表
        name_entry.bind("<KeyRelease>", self.update_person_data)
        id_entry.bind("<KeyRelease>", self.update_person_data)
        
        # 添加身份证号校验逻辑 - 使用FocusOut事件避免每次按键都校验
        id_entry.bind("<FocusOut>", lambda event, entry=id_entry, var=id_var: self.validate_id_card(var, entry))
        
        # 添加身份证号码变化时的重置验证状态
        id_var.trace_add("write", lambda *args, entry=id_entry: self.reset_id_validation(entry))
        
        name_var.trace_add("write", lambda *args: self.update_person_data())
        id_var.trace_add("write", lambda *args: self.update_person_data())
        
        # 绑定回车键，像Excel一样导航到下一行同列位置
        name_entry.bind("<Return>", lambda event: self._focus_next_row_name(row_idx))
        id_entry.bind("<Return>", lambda event: self._focus_next_row_id(row_idx))
        
        self.member_name_vars.append(name_var)
        self.member_id_vars.append(id_var)
        self.member_frames.append(member_frame)
        
        # 初始化标记，用于跟踪是否已经提示过错误
        if not hasattr(self, 'id_error_shown'):
            self.id_error_shown = set()
        
        return member_frame
    
    def add_family_member(self):
        """添加家庭成员"""
        new_frame = self.add_family_member_row()
        self.update_person_data()
        
        # 由于我们已经在add_family_member_row中应用了正确的尺寸，
        # 这里不需要额外的缩放处理
        # 确保后续不会被其他缩放处理覆盖
        self.root.update()  # 立即更新UI确保尺寸应用生效
    
    def delete_family_member(self, frame, idx):
        """删除家庭成员"""
        # 至少保留一个家庭成员
        if len(self.member_frames) <= 1:
            messagebox.showinfo("提示", "至少需要保留一个家庭成员")
            return
        
        # 删除UI元素
        frame.destroy()
        
        # 从列表中移除对应的变量和控件
        del self.member_name_vars[idx]
        del self.member_id_vars[idx]
        del self.member_entries[idx]
        del self.member_frames[idx]
        
        # 更新人员数据
        self.update_person_data()
        
        # 重新排列剩余成员的索引
        for i, mframe in enumerate(self.member_frames):
            for child in mframe.winfo_children():
                if isinstance(child, ttk.Button) and child['text'] == "删除":
                    child.configure(command=lambda frame=mframe, idx=i: self.delete_family_member(frame, idx))
    
    def update_person_data(self, event=None):
        """更新人员数据列表，用于产权人选择"""
        self.person_data = []
        
        # 添加所有家庭成员信息
        for i in range(len(self.member_name_vars)):
            name = self.member_name_vars[i].get().strip()
            id_card = self.member_id_vars[i].get().strip()
            if name:
                self.person_data.append((name, id_card))
        
        # 更新家庭成员总人数 - 但不显示在UI上
        # self.members_count_label.config(text=f"({len(self.person_data)}人)")
        
        # 更新所有房产表单中的下拉菜单选项
        self.update_all_apartment_comboboxes()
        
        # 更新房产总套数 - 但不显示在UI上
        # self.update_apartment_count()
    
    def update_apartment_count(self):
        """更新房产总套数"""
        valid_apartments = 0
        for apt in self.apartment_frames:
            # 只要有一个所有者，就认为是有效房产
            if any(var.get().strip() for var in apt["owner_vars"]):
                valid_apartments += 1
        
        # 不显示总套数
        # self.apartments_count_label.config(text=f"(共{valid_apartments}套)")
    
    def update_all_apartment_comboboxes(self):
        """更新所有房产中的人员选择下拉菜单"""
        names = [""] + [person[0] for person in self.person_data]
        for apt in self.apartment_frames:
            # 更新所有所有者下拉菜单
            for i, combo in enumerate(apt["owner_combos"]):
                current_value = apt["owner_vars"][i].get()
                
                # 设置新的下拉菜单值
                combo["values"] = names
                
                # 如果当前值在新名单中，保持选择
                if current_value in names:
                    apt["owner_vars"][i].set(current_value)
                # 否则清除选择
                elif current_value and current_value not in names:
                    apt["owner_vars"][i].set("")
    
    def on_owner_selected(self, event, apt_data, owner_idx):
        """当产权人被选择时更新身份证号"""
        selected_name = apt_data["owner_vars"][owner_idx].get()
        
        # 根据选择的姓名找到对应的身份证号
        if selected_name:
            for name, id_card in self.person_data:
                if name == selected_name:
                    apt_data["id_card_vars"][owner_idx].set(id_card)
                    break
        else:
            # 如果没有选择名字，清空身份证号
            apt_data["id_card_vars"][owner_idx].set("")
    
    def add_apartment_frame(self):
        """添加一个新的公寓信息输入框架，使用表格形式 - 更符合图片样式"""
        # 创建行容器 - 第一行时创建表格边框
        if not self.apartment_frames:
            # 创建表格边框 - 第一次添加行时
            table_border = ttk.Frame(self.apartments_container, style="Table.TFrame", relief=tk.RIDGE, borderwidth=1)
            table_border.pack(fill=tk.X)
            row_container = ttk.Frame(table_border)
            row_container.pack(fill=tk.X)
        else:
            # 已有表格，直接在内部添加行
            row_container = self.apartment_frames[0]["container"] if self.apartment_frames else self.apartments_container
        
        # 创建行容器
        row_frame = ttk.Frame(row_container, style="Table.TFrame")
        row_frame.pack(fill=tk.X)
        
        # 设置输入框统一样式
        style = ttk.Style()
        style.configure('Table.TEntry', font=('黑体', 14))  # 增大字体
        style.configure('Table.TCombobox', font=('黑体', 14))  # 增大字体
        style.configure('Table.TButton', font=('黑体', 14))  # 增大字体
        style.configure('Tall.TCombobox', padding=(5, 10), font=('黑体', 14))  # 增加下拉框高度与输入框一致
        style.configure('AddOwner.TButton', font=('黑体', 14, 'bold'), padding=(0, 10))  # 增加按钮高度与输入框一致
        
        # 使用保存的列宽
        widths = self.column_widths
        
        # 创建每个单元格 - 使用水平分隔线而不是完整边框
        # 楼栋号
        building_var = tk.StringVar()
        building_entry = ttk.Entry(row_frame, textvariable=building_var, width=widths[0], 
                                  justify=tk.CENTER, font=('黑体', 14), style='Tall.TEntry')  # 增加输入框高度
        building_entry.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        
        # 单元号
        unit_var = tk.StringVar()
        unit_entry = ttk.Entry(row_frame, textvariable=unit_var, width=widths[1], 
                              justify=tk.CENTER, font=('黑体', 14), style='Tall.TEntry')  # 增加输入框高度
        unit_entry.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        
        # 房间号
        room_var = tk.StringVar()
        room_entry = ttk.Entry(row_frame, textvariable=room_var, width=widths[2], 
                              justify=tk.CENTER, font=('黑体', 14), style='Tall.TEntry')  # 增加输入框高度
        room_entry.grid(row=0, column=2, sticky="nsew", padx=0, pady=0)
        
        # 面积
        area_var = tk.StringVar()
        area_entry = ttk.Entry(row_frame, textvariable=area_var, width=widths[3], 
                              justify=tk.CENTER, font=('黑体', 14), style='Tall.TEntry')  # 增加输入框高度
        area_entry.grid(row=0, column=3, sticky="nsew", padx=0, pady=0)
        
        # 添加所有者容器 - 支持多个所有者
        owners_container = ttk.Frame(row_frame)
        owners_container.grid(row=0, column=4, columnspan=3, sticky="nsew", padx=0, pady=0)
        
        # 存储所有者变量列表
        owner_vars = []
        id_card_vars = []
        owner_combos = []
        
        # 添加第一个默认所有者
        owner_var = tk.StringVar()
        id_card_var = tk.StringVar()
        
        # 产权人
        names = [""] + [person[0] for person in self.person_data]
        owner_combo = ttk.Combobox(owners_container, textvariable=owner_var, values=names, 
                                   width=widths[4], font=('黑体', 14), style='Tall.TCombobox')  # 增加高度
        owner_combo.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        
        # "+"按钮 - 只在第一行显示，设置takefocus=0使Tab键跳过此按钮
        add_btn = ttk.Button(owners_container, text="+", width=widths[5], style='AddOwner.TButton', takefocus=0)
        add_btn.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        
        # 身份证号 - 确保宽度合适
        id_card_width = widths[6] - 2  # 减少宽度以适应布局
        id_entry = ttk.Entry(owners_container, textvariable=id_card_var, width=id_card_width, 
                            font=('黑体', 14), style='Tall.TEntry')  # 增加输入框高度
        id_entry.grid(row=0, column=3, sticky="nsew", padx=0, pady=0)
        
        # 存储第一个所有者
        owner_vars.append(owner_var)
        id_card_vars.append(id_card_var)
        owner_combos.append(owner_combo)
        
        # 删除按钮 - 使用与家庭成员部分一致的样式，设置takefocus=0使Tab键跳过此按钮
        delete_btn = ttk.Button(row_frame, text="删除", width=widths[7], style='Large.TButton', takefocus=0)
        delete_btn.grid(row=0, column=7, sticky="nsew", padx=0, pady=0)
        
        # 确保列宽一致
        for i in range(8):
            row_frame.columnconfigure(i, weight=1)
        
        # 获取当前公寓索引，用于回车键导航
        apt_idx = len(self.apartment_frames)
        
        # 存储UI变量的字典 - 使用列表存储多个所有者
        apartment_data = {
            "frame": row_frame,
            "owners_container": owners_container,
            "container": row_container if not self.apartment_frames else self.apartment_frames[0]["container"],
            "building": building_var,
            "unit": unit_var,
            "room": room_var,
            "area": area_var,
            "owner_vars": owner_vars,
            "id_card_vars": id_card_vars,
            "owner_combos": owner_combos,
            "owner_rows": [0],  # 记录每个所有者在container中的行号
            "add_btn": add_btn,
            "delete_btn": delete_btn,
            "building_entry": building_entry,
            "unit_entry": unit_entry,
            "room_entry": room_entry,
            "area_entry": area_entry,
            "id_entries": [id_entry],
            "apt_idx": apt_idx  # 记录公寓索引，用于回车键导航
        }
        
        # 绑定回车键事件 - 像Excel一样垂直导航
        building_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, "building_entry"))
        unit_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, "unit_entry"))
        room_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, "room_entry"))
        area_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, "area_entry"))
        owner_combo.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, 0))
        id_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, 0, is_id=True))
        
        # 绑定选择事件
        owner_combo.bind("<<ComboboxSelected>>", lambda event, apt=apartment_data, idx=0: 
                         self.on_owner_selected(event, apt, idx))
        
        # 绑定添加所有者按钮事件
        add_btn.configure(command=lambda apt=apartment_data: self.add_owner_to_apartment(apt))
        
        # 配置删除按钮
        delete_btn.configure(command=lambda apt=apartment_data: self.delete_apartment(apt))
        
        # 应用当前缩放比例到新添加的控件
        if hasattr(self, 'scale_factor') and self.scale_factor != 1.0:
            for child in row_frame.winfo_children():
                self._update_widget_fonts_recursive(child, self.scale_factor)
                self._update_widget_sizes_recursive(child, self.scale_factor)
            
            for child in apartment_data["owners_container"].winfo_children():
                self._update_widget_fonts_recursive(child, self.scale_factor)
                self._update_widget_sizes_recursive(child, self.scale_factor)
        
        self.apartment_frames.append(apartment_data)
        
        # 更新房产总套数
        self.update_apartment_count()
        
        return apartment_data
    
    def add_owner_to_apartment(self, apartment_data):
        """为指定的公寓添加额外的所有者"""
        # 获取当前所有者的行号并计算新所有者的行号
        next_row = len(apartment_data["owner_rows"])
        
        # 创建新变量
        owner_var = tk.StringVar()
        id_card_var = tk.StringVar()
        
        # 创建所有者UI
        owners_container = apartment_data["owners_container"]
        widths = self.column_widths
        
        # 基础控件尺寸
        owner_width = widths[4]
        id_card_width = widths[6] - 2  # 减少宽度以适应布局
        
        # 如果有缩放因子，计算实际应用的尺寸和字体大小
        applied_owner_width = owner_width
        applied_id_card_width = id_card_width
        applied_font_size = 14
        
        if hasattr(self, 'scale_factor') and self.scale_factor != 1.0:
            applied_owner_width = int(owner_width * self.scale_factor)
            applied_id_card_width = int(id_card_width * self.scale_factor)
            applied_font_size = int(14 * self.scale_factor)
        
        # 产权人下拉框 - 直接应用正确尺寸
        names = [""] + [person[0] for person in self.person_data]
        owner_combo = ttk.Combobox(owners_container, textvariable=owner_var, values=names, 
                                  width=applied_owner_width, font=('黑体', applied_font_size), 
                                  style='Tall.TCombobox')
        owner_combo.grid(row=next_row, column=0, sticky="nsew", padx=0, pady=0)
        
        # 新增行只有"-"按钮，没有"+"按钮，设置takefocus=0使Tab键跳过此按钮
        # +/-按钮保持固定尺寸
        remove_btn = ttk.Button(owners_container, text="-", width=widths[5], style='AddOwner.TButton', takefocus=0)
        remove_btn.grid(row=next_row, column=1, sticky="nsew", padx=0, pady=0)
        
        # 身份证号 - 直接应用正确尺寸
        id_entry = ttk.Entry(owners_container, textvariable=id_card_var, width=applied_id_card_width, 
                           font=('黑体', applied_font_size), style='Tall.TEntry')
        id_entry.grid(row=next_row, column=3, sticky="nsew", padx=0, pady=0)
        
        # 存储变量
        apartment_data["owner_vars"].append(owner_var)
        apartment_data["id_card_vars"].append(id_card_var)
        apartment_data["owner_combos"].append(owner_combo)
        apartment_data["owner_rows"].append(next_row)
        apartment_data["id_entries"].append(id_entry)
        
        # 为"-"按钮绑定删除事件
        remove_btn.configure(command=lambda apt=apartment_data, row=next_row: 
                              self.remove_owner_from_apartment(apt, row))
        
        # 绑定选择事件
        owner_combo.bind("<<ComboboxSelected>>", lambda event, apt=apartment_data, idx=next_row: 
                         self.on_owner_selected(event, apt, idx))
        
        # 绑定回车键事件，垂直导航
        apt_idx = apartment_data.get("apt_idx", 0)
        owner_combo.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, next_row))
        id_entry.bind("<Return>", lambda event: self._focus_down_cell(apt_idx, next_row, is_id=True))
        
        # 存储原始控件尺寸信息，用于缩放计算
        if hasattr(self, 'original_widget_sizes'):
            self.original_widget_sizes[str(owner_combo)] = owner_width
            self.original_widget_sizes[str(id_entry)] = id_card_width
        
        # 聚焦到新添加的行中的产权人下拉框
        owner_combo.focus_set()
        
        # 刷新界面
        self.root.update()
    
    def remove_owner_from_apartment(self, apartment_data, row_idx):
        """从指定的公寓中删除所有者行"""
        # 阻止删除最后一个所有者
        if len(apartment_data["owner_vars"]) <= 1:
            messagebox.showinfo("提示", "每个房产至少需要一个所有者")
            return
        
        # 获取行索引
        idx = apartment_data["owner_rows"].index(row_idx)
        
        # 移除UI组件
        for widget in apartment_data["owners_container"].grid_slaves():
            if int(widget.grid_info()["row"]) == row_idx:
                widget.destroy()
        
        # 从数据中移除
        del apartment_data["owner_vars"][idx]
        del apartment_data["id_card_vars"][idx]
        del apartment_data["owner_combos"][idx]
        del apartment_data["owner_rows"][idx]
        
        # 刷新界面
        self.root.update()
    
    def delete_apartment(self, apartment_data):
        """删除一个公寓表单"""
        if len(self.apartment_frames) <= 1:
            messagebox.showinfo("提示", "至少需要保留一个房产信息")
            return
        
        self.apartment_frames.remove(apartment_data)
        apartment_data["frame"].destroy()
        
        # 更新房产总套数
        self.update_apartment_count()
    
    def clear_form(self, confirm=True):
        """清空所有表单数据"""
        if confirm and not messagebox.askyesno("确认", "确定要清空所有数据吗？"):
            return
        
        # 清空家庭成员信息
        for var in self.member_name_vars + self.member_id_vars:
            var.set("")
        
        # 清空项目信息
        self.project_var.set("")
        self.community_var.set("")
        
        # 清空所有房产信息（除了第一个）
        while len(self.apartment_frames) > 1:
            apt = self.apartment_frames[-1]
            self.delete_apartment(apt)
        
        # 清空第一个房产信息
        if self.apartment_frames:
            first_apt = self.apartment_frames[0]
            first_apt["building"].set("")
            first_apt["unit"].set("")
            first_apt["room"].set("")
            first_apt["area"].set("")
            for i in range(len(first_apt["owner_vars"])):
                first_apt["owner_vars"][i].set("")
            first_apt["id_card_vars"][0].set("")
        
        # 更新人员数据和统计
        self.update_person_data()
        
        self.status_var.set("表单已清空")
    
    def collect_data(self):
        """收集表单中的所有数据"""
        # 收集家庭成员信息
        members = []
        for i in range(len(self.member_name_vars)):
            name = self.member_name_vars[i].get().strip()
            id_card = self.member_id_vars[i].get().strip()
            if name:
                members.append({"name": name, "id_card": id_card})
        
        # 收集基本信息
        data = {
            "members": members,
            "project": self.project_var.get().strip(),
            "community": self.community_var.get().strip(),
            "members_count": len(members),
            "apartments": []
        }
        
        # 收集房产信息
        for apt in self.apartment_frames:
            # 收集所有有效的所有者
            owners = []
            for i in range(len(apt["owner_vars"])):
                owner_name = apt["owner_vars"][i].get().strip()
                id_card = apt["id_card_vars"][i].get().strip()
                if owner_name:
                    owners.append({"name": owner_name, "id_card": id_card})
            
            apartment = {
                "building": apt["building"].get().strip(),
                "unit": apt["unit"].get().strip(),
                "room": apt["room"].get().strip(),
                "area": apt["area"].get().strip(),
                "owners": owners
            }
            
            # 只添加有所有者的公寓
            if apartment["owners"]:
                data["apartments"].append(apartment)
        
        data["apartments_count"] = len(data["apartments"])
        
        return data
    
    def validate_data(self, data):
        """验证数据的完整性"""
        if not data["members"]:
            return "请至少添加一位家庭成员"
        
        if not data["project"]:
            return "项目名称不能为空"
        
        if not data["community"]:
            return "小区名称不能为空"
        
        if not data["apartments"]:
            return "至少需要添加一个有效的房产信息"
        
        for apt in data["apartments"]:
            if not apt["building"]:
                return "楼栋号不能为空"
            if not apt["unit"]:
                return "单元号不能为空"
            if not apt["room"]:
                return "房间号不能为空"
            if not apt["area"]:
                return "建筑面积不能为空"
            if not apt["owners"]:
                return "产权人不能为空"
            
            for owner in apt["owners"]:
                if not owner["id_card"]:
                    return f"产权人 {owner['name']} 的身份证号不能为空"
        
        return None  # 验证通过
    
    def generate_document_text_vba_style(self, data):
        """生成文档的文本内容，与VBA代码格式一致"""
        text = "情况说明\n\n"
        text += "市不动产登记中心：\n"
        
        # 使用与VBA相同的格式
        main_person = data["members"][0] if data["members"] else {"name": "", "id_card": ""}
        main_apt = data["apartments"][0] if data["apartments"] else {"building": "", "unit": "", "room": "", "area": ""}
        main_owner = main_apt["owners"][0] if main_apt.get("owners") else {"name": "", "id_card": ""}
        
        text += f"兹有我辖区{data['project']}项目拆迁户{main_person['name']}（身份证号{main_person['id_card']}）,"
        text += f"安置于{data['community']}小区:{main_apt['building']}栋{main_apt['unit']}单元{main_apt['room']}室,"
        text += f"建筑面积{main_apt['area']}㎡"
        
        # 添加额外的公寓信息
        if len(data["apartments"]) > 1:
            for i in range(1, len(data["apartments"])):
                apt = data["apartments"][i]
                text += f"、{apt['building']}栋{apt['unit']}单元{apt['room']}室,建筑面积{apt['area']}㎡"
        
        text += "。\n"
        
        # 添加产权信息
        text += f"因政策调整及根据分配协议，将《{data['community']}小区还原安置不动产登记明细表》中的产权人姓名、产权人身份证号更正如下:\n"
        
        # 输出每套房产的产权人信息，支持多个产权人
        for apt in data["apartments"]:
            text += f"{apt['building']}栋{apt['unit']}单元{apt['room']}室,建筑面积{apt['area']}㎡"
            
            # 添加所有产权人
            for i, owner in enumerate(apt["owners"]):
                if i == 0:
                    text += f"为{owner['name']}，身份证号为{owner['id_card']}"
                else:
                    text += f"；{owner['name']}，身份证号为{owner['id_card']}"
            
            text += "。\n"
        
        text += "特此说明。\n\n"
        text += f"{data['community']}安置办\n"
        
        # 添加当前日期，格式为YYYY年MM月DD日
        now = datetime.now()
        text += now.strftime("%Y年%m月%d日")
        
        return text
    
    def generate_word_document(self, data):
        """生成WPS兼容的Word文档，完全按照VBA代码的格式设置"""
        try:
            # 创建新的Word文档
            doc = Document()
            
            # 设置中文字体 - 整体文档使用宋体三号
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = Pt(16)  # 宋体三号
            doc.styles['Normal'].font.bold = False  # 不加粗
            
            # 先添加一个空行 - 整体下移
            first_empty_para = doc.add_paragraph()
            first_empty_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            first_empty_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            first_empty_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            first_empty_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            
            # 添加标题
            title_para = doc.add_paragraph("情况说明")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            # WPS 具体的行距设置
            title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            title_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            title_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            title_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            for run in title_para.runs:
                run.font.size = Pt(22)  # 标题字号
                run.font.bold = True
            
            # 只添加一个空行 - 减少标题和正文之间的空间
            empty_para = doc.add_paragraph()
            empty_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            empty_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            empty_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            empty_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            
            # 添加"市不动产登记中心："，并加粗
            reg_para = doc.add_paragraph()
            reg_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
            reg_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            reg_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            reg_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            reg_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            reg_run = reg_para.add_run("市不动产登记中心：")
            reg_run.font.bold = True
            
            # 主要内容
            main_person = data["members"][0] if data["members"] else {"name": "", "id_card": ""}
            main_apt = data["apartments"][0] if data["apartments"] else {"building": "", "unit": "", "room": "", "area": ""}
            main_owner = main_apt["owners"][0] if main_apt.get("owners") else {"name": "", "id_card": ""}
            
            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
            content_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进2字符（WPS中接近0.5英寸）
            content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            content_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            content_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            content_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            
            content_para.add_run(f"兹有我辖区{data['project']}项目拆迁户{main_person['name']}（身份证号{main_person['id_card']}）,")
            content_para.add_run(f"安置于{data['community']}小区:{main_apt['building']}栋{main_apt['unit']}单元{main_apt['room']}室,")
            content_para.add_run(f"建筑面积{main_apt['area']}㎡")
            
            # 添加额外的公寓信息
            if len(data["apartments"]) > 1:
                for i in range(1, len(data["apartments"])):
                    apt = data["apartments"][i]
                    content_para.add_run(f"、{apt['building']}栋{apt['unit']}单元{apt['room']}室,建筑面积{apt['area']}㎡")
            
            content_para.add_run("。")
            
            # 添加产权调整说明
            adjust_para = doc.add_paragraph()
            adjust_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
            adjust_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进2字符（WPS中接近0.5英寸）
            adjust_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            adjust_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            adjust_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            adjust_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            adjust_para.add_run(f"因政策调整及根据分配协议，将《{data['community']}小区还原安置不动产登记明细表》中的产权人姓名、产权人身份证号更正如下:")
            
            # 添加每套房产的产权人信息，支持多个产权人
            for apt in data["apartments"]:
                apt_para = doc.add_paragraph()
                apt_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
                apt_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进2字符（WPS中接近0.5英寸）
                apt_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                apt_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
                apt_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
                apt_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
                
                apt_para.add_run(f"{apt['building']}栋{apt['unit']}单元{apt['room']}室,建筑面积{apt['area']}㎡")
                
                # 添加所有产权人
                for i, owner in enumerate(apt["owners"]):
                    if i == 0:
                        apt_para.add_run(f"为{owner['name']}，身份证号为{owner['id_card']}")
                    else:
                        apt_para.add_run(f"；{owner['name']}，身份证号为{owner['id_card']}")
                
                apt_para.add_run("。")
            
            # 添加结尾
            end_para = doc.add_paragraph()
            end_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
            end_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进2字符（WPS中接近0.5英寸）
            end_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            end_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            end_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            end_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            end_para.add_run("特此说明。")
            
            # 添加空行
            empty_para2 = doc.add_paragraph()
            empty_para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            empty_para2.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            empty_para2.paragraph_format.space_after = Pt(0)  # 段后间距为0
            empty_para2.paragraph_format.space_before = Pt(0)  # 段前间距为0
            
            # 添加落款
            sign_para = doc.add_paragraph()
            sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
            sign_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            sign_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            sign_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            sign_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            sign_para.add_run(f"{data['community']}安置办")
            
            # 添加日期
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
            date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            date_para.paragraph_format.line_spacing = Pt(30)  # 固定行距30磅
            date_para.paragraph_format.space_after = Pt(0)  # 段后间距为0
            date_para.paragraph_format.space_before = Pt(0)  # 段前间距为0
            now = datetime.now()
            # 修改日期格式，去掉月份前导零
            formatted_date = f"{now.year}年{now.month}月{now.day}日"
            date_para.add_run(formatted_date)
            
            # 设置页面尺寸为A4纸和页面边距
            from docx.shared import Mm
            sections = doc.sections
            for section in sections:
                # 设置纸张大小为A4（210mm × 297mm）
                section.page_width = Mm(210)
                section.page_height = Mm(297)
                
                # 设置页面边距
                section.top_margin = Pt(72)     # 保持上下边距为1英寸
                section.bottom_margin = Pt(72)  # 保持上下边距为1英寸
                section.left_margin = Pt(90)    # 左边距改为3.18厘米
                section.right_margin = Pt(90)   # 右边距改为3.18厘米
            
            return doc
        except Exception as e:
            raise Exception(f"生成Word文档时出错: {str(e)}")
    
    def generate_report(self):
        try:
            # 收集数据
            data = self.collect_data()
            
            # 验证数据
            error = self.validate_data(data)
            if error:
                messagebox.showerror("错误", error)
                return
            
            self.status_var.set("正在生成情况说明...")
            self.root.update()
            
            # 生成Word文档
            doc = self.generate_word_document(data)
            
            # 设置Word文档保存目录 - 使用用户Documents目录
            save_dir = self.doc_dir
            
            # 获取第一个成员的姓名作为默认文件名
            first_member = data["members"][0]["name"] if data["members"] else "安置户"
            default_filename = f"{first_member}户安置房分配协议.docx"
            
            # 创建默认的数据文件名
            data_filename = f"{first_member}户安置房分配协议数据.dat"
            
            # 设置完整的文件路径
            save_path = os.path.join(save_dir, default_filename)
            data_path = os.path.join(self.config_dir, data_filename)  # 数据文件保存到CONFIG目录
            
            # 保存表单数据
            self.save_data_to_file(data_path)
            
            # 直接保存文档到默认位置
            doc.save(save_path)
            
            # 直接打开文档，不询问
            os.startfile(save_path)
            
            self.status_var.set(f"文档已保存并打开: {save_path}")
        
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.status_var.set("生成失败")
    
    def show_about(self):
        messagebox.showinfo("关于", "集体土地安置房分配协议生成器 v1.0\n\n本程序用于直接输入数据生成Word格式的分配协议文档。")

    def _focus_down_cell(self, apt_idx, field_idx, is_id=False):
        """按下回车时导航到下一行的相同位置，像Excel一样的垂直导航"""
        # 如果当前公寓索引超出范围，返回
        if apt_idx >= len(self.apartment_frames):
            return
        
        current_apt = self.apartment_frames[apt_idx]
        
        # 如果是主要字段（building, unit, room, area）
        if isinstance(field_idx, str):
            # 检查下一个公寓是否存在
            if apt_idx + 1 < len(self.apartment_frames):
                next_apt = self.apartment_frames[apt_idx + 1]
                # 导航到下一个公寓的同名字段
                if field_idx in next_apt:
                    next_apt[field_idx].focus_set()
                else:
                    # 如果没有找到相同字段，尝试导航到第一个字段
                    next_apt["building_entry"].focus_set()
            else:
                # 没有下一个公寓，添加新公寓
                pass
        else:
            # 如果是所有者字段（owner_combo 或 id_entry）
            if is_id:  # 如果是身份证号字段
                # 检查当前公寓中是否有下一行
                if field_idx + 1 < len(current_apt["id_entries"]):
                    # 导航到当前公寓的下一行同列
                    current_apt["id_entries"][field_idx + 1].focus_set()
                elif apt_idx + 1 < len(self.apartment_frames):
                    # 导航到下一个公寓的第一行
                    next_apt = self.apartment_frames[apt_idx + 1]
                    if next_apt["id_entries"]:
                        next_apt["id_entries"][0].focus_set()
                    else:
                        next_apt["building_entry"].focus_set()
            else:  # 如果是产权人下拉框
                # 检查当前公寓中是否有下一行
                if field_idx + 1 < len(current_apt["owner_combos"]):
                    # 导航到当前公寓的下一行同列
                    current_apt["owner_combos"][field_idx + 1].focus_set()
                elif apt_idx + 1 < len(self.apartment_frames):
                    # 导航到下一个公寓的第一行
                    next_apt = self.apartment_frames[apt_idx + 1]
                    if next_apt["owner_combos"]:
                        next_apt["owner_combos"][0].focus_set()
                    else:
                        next_apt["building_entry"].focus_set()

    def _focus_next_row_name(self, current_row):
        """在姓名输入框之间垂直导航"""
        # 检查是否有下一行
        if current_row + 1 < len(self.member_entries):
            # 导航到下一行的姓名输入框
            self.member_entries[current_row + 1][0].focus_set()
        else:
            # 这是最后一行，可以考虑添加新行或转到其他区域
            # 例如，导航到项目信息区域或第一个公寓
            if self.apartment_frames:
                self.apartment_frames[0]["building_entry"].focus_set()

    def _focus_next_row_id(self, current_row):
        """在身份证号输入框之间垂直导航"""
        # 检查是否有下一行
        if current_row + 1 < len(self.member_entries):
            # 导航到下一行的身份证号输入框
            self.member_entries[current_row + 1][1].focus_set()
        else:
            # 这是最后一行，可以考虑添加新行或转到其他区域
            if self.apartment_frames:
                self.apartment_frames[0]["building_entry"].focus_set()

    def _configure_scroll_region(self, event=None):
        """配置画布的滚动区域"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def _configure_canvas(self, event):
        """根据画布大小调整内部框架宽度"""
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas_frame, width=canvas_width)
    
    def print_report(self):
        """直接保存并打印情况说明，不打开Word文档"""
        # 创建加载对话框
        loading_dialog = LoadingDialog(self.root, title="打印中", message="正在生成并打印文档，请稍候...")
        
        # 定义打印任务的函数
        def print_task():
            try:
                # 初始化COM环境，解决线程中使用COM的问题
                pythoncom.CoInitialize()
                
                # 收集数据
                data = self.collect_data()
                
                # 验证数据
                error = self.validate_data(data)
                if error:
                    # 关闭对话框并显示错误
                    self.root.after(0, lambda: loading_dialog.stop())
                    self.root.after(100, lambda err=error: messagebox.showerror("错误", err))
                    return
                
                # 更新状态和对话框消息
                self.status_var.set("正在生成文档...")
                self.root.after(0, lambda: loading_dialog.update_message("正在生成文档..."))
                
                # 生成Word文档
                doc = self.generate_word_document(data)
                
                # 设置Word文档保存目录 - 使用用户Documents目录
                save_dir = self.doc_dir
                
                # 获取第一个成员的姓名作为默认文件名
                first_member = data["members"][0]["name"] if data["members"] else "安置户"
                default_filename = f"{first_member}户安置房分配协议.docx"
                
                # 创建默认的数据文件名
                data_filename = f"{first_member}户安置房分配协议数据.dat"
                
                # 设置完整的文件路径
                save_path = os.path.join(save_dir, default_filename)
                data_path = os.path.join(self.config_dir, data_filename)  # 数据文件保存到CONFIG目录
                
                # 更新对话框消息
                self.root.after(0, lambda: loading_dialog.update_message("正在保存文档..."))
                
                # 保存表单数据
                self.save_data_to_file(data_path)
                
                # 保存文档到默认位置
                doc.save(save_path)
                
                # 更新对话框消息
                self.root.after(0, lambda: loading_dialog.update_message("正在发送到打印机..."))
                
                # 静默打印文档
                # 使用COM接口操作Word，实现真正的静默打印
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False  # 不显示Word窗口
                
                # 打开文档
                doc = word.Documents.Open(save_path)
                
                # 直接打印
                doc.PrintOut()
                
                # 关闭文档和Word应用程序
                doc.Close(SaveChanges=False)
                word.Quit()
                
                # 更新状态栏
                self.status_var.set(f"文档已保存并发送到打印机: {save_path}")
                
                # 关闭对话框并显示成功消息
                self.root.after(0, lambda: loading_dialog.stop())
                self.root.after(100, lambda: messagebox.showinfo("打印", "文档已发送到默认打印机。"))
                
            except Exception as e:
                # 捕获异常消息
                error_message = str(e)
                # 关闭对话框并显示错误
                self.root.after(0, lambda: loading_dialog.stop())
                self.root.after(100, lambda err=error_message: messagebox.showerror("错误", err))
                self.status_var.set("打印失败")
            finally:
                # 释放COM环境
                pythoncom.CoUninitialize()
        
        # 创建并启动打印线程
        print_thread = threading.Thread(target=print_task)
        print_thread.daemon = True  # 设置为守护线程，这样主程序退出时线程也会退出
        print_thread.start()

    def open_documents_folder(self):
        """打开安置房分配协议文档所在文件夹"""
        try:
            # 确保文件夹存在
            if not os.path.exists(self.doc_dir):
                os.makedirs(self.doc_dir)
                
            # 打开文件夹
            if sys.platform == 'win32':
                os.startfile(self.doc_dir)
            elif sys.platform == 'darwin':  # macOS
                subprocess.call(['open', self.doc_dir])
            else:  # Linux
                subprocess.call(['xdg-open', self.doc_dir])
                
            self.status_var.set(f"已打开文档文件夹: {self.doc_dir}")
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文档文件夹: {str(e)}")
            self.status_var.set("打开文档文件夹失败")

    def open_recent_document(self):
        """打开最近的文档文件"""
        try:
            # 确保文档目录存在
            if not os.path.exists(self.doc_dir):
                os.makedirs(self.doc_dir)
                messagebox.showinfo("提示", "文档文件夹为空，没有可打开的文档。")
                return
                
            # 获取所有docx文件
            docx_files = [f for f in os.listdir(self.doc_dir) if f.endswith('.docx')]
            
            if not docx_files:
                messagebox.showinfo("提示", "文档文件夹中没有找到Word文档。")
                return
                
            # 按修改时间排序，获取最新的文档
            docx_files.sort(key=lambda x: os.path.getmtime(os.path.join(self.doc_dir, x)), reverse=True)
            
            # 选择最新的文档打开
            latest_doc = os.path.join(self.doc_dir, docx_files[0])
            
            # 打开文档
            if sys.platform == 'win32':
                os.startfile(latest_doc)
            elif sys.platform == 'darwin':  # macOS
                subprocess.call(['open', latest_doc])
            else:  # Linux
                subprocess.call(['xdg-open', latest_doc])
                
            self.status_var.set(f"已打开最新文档: {docx_files[0]}")
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文档: {str(e)}")
            self.status_var.set("打开文档失败")
            
            # 如果打开文档失败，尝试打开文档文件夹
            try:
                self.open_documents_folder()
            except:
                pass

    def reset_id_validation(self, entry):
        """当身份证号码输入框内容改变时，重置验证状态"""
        entry_id = str(entry)
        if entry_id in self.id_error_shown:
            # 从已提示集合中移除，允许再次验证
            self.id_error_shown.remove(entry_id)
    
    def validate_id_card(self, id_var, entry):
        """验证身份证号码的有效性"""
        id_card = id_var.get().strip()
        
        # 如果字段为空，不验证
        if not id_card:
            return True
            
        # 获取输入框的唯一标识
        entry_id = str(entry)
        
        # 验证身份证号码
        result, message = self.check_id_card(id_card)
        
        # 如果验证失败，显示错误消息
        if not result:
            # 检查该输入框是否已经提示过错误
            if entry_id in self.id_error_shown:
                # 已经提示过，但是内容已经改变，所以需要重新验证
                # 从集合中移除，允许再次提示
                self.id_error_shown.remove(entry_id)
                
            # 如果不在已提示集合中，显示错误消息
            if entry_id not in self.id_error_shown:
                messagebox.showerror("身份证号码错误", message)
                self.id_error_shown.add(entry_id)  # 标记已经提示过错误
            return False
        else:
            # 验证通过，如果之前有错误，现在移除错误标记
            if entry_id in self.id_error_shown:
                self.id_error_shown.remove(entry_id)
            return True
    
    def check_id_card(self, id_card):
        """检查身份证号码的有效性，返回(是否有效, 错误信息)"""
        id_card = id_card.upper()  # 转换X为大写
        
        # 基本格式验证
        if len(id_card) == 18:
            # 18位身份证验证
            if not (id_card[:-1].isdigit() and (id_card[-1].isdigit() or id_card[-1] == 'X')):
                return False, "18位身份证前17位应为数字，最后一位应为数字或X"
                
            # 验证出生日期
            try:
                birth_year = int(id_card[6:10])
                birth_month = int(id_card[10:12])
                birth_day = int(id_card[12:14])
                
                # 简单的日期检查
                if birth_year < 1900 or birth_year > datetime.now().year:
                    return False, f"出生年份 {birth_year} 不合理"
                if birth_month < 1 or birth_month > 12:
                    return False, f"出生月份 {birth_month} 不合理"
                if birth_day < 1 or birth_day > 31:
                    return False, f"出生日期 {birth_day} 不合理"
                    
                # 检查2月份日期
                if birth_month == 2:
                    is_leap = (birth_year % 4 == 0 and birth_year % 100 != 0) or (birth_year % 400 == 0)
                    if (is_leap and birth_day > 29) or (not is_leap and birth_day > 28):
                        return False, f"{birth_year}年2月没有{birth_day}日"
                        
                # 检查小月
                if birth_month in [4, 6, 9, 11] and birth_day > 30:
                    return False, f"{birth_month}月没有{birth_day}日"
            except:
                return False, "出生日期格式不正确"
                
            # 验证校验码（第18位）
            factor = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
            parity = ['1', '0', 'X', '9', '8', '7', '6', '5', '4', '3', '2']
            
            sum = 0
            for i in range(17):
                sum += int(id_card[i]) * factor[i]
                
            if parity[sum % 11] != id_card[17]:
                return False, "身份证校验码不正确"
                
            return True, ""
            
        elif len(id_card) == 15:
            # 15位身份证验证 (旧版)
            if not id_card.isdigit():
                return False, "15位身份证应全为数字"
                
            # 验证出生日期
            try:
                birth_year = int('19' + id_card[6:8])  # 15位身份证出生年份为19XX年
                birth_month = int(id_card[8:10])
                birth_day = int(id_card[10:12])
                
                # 简单的日期检查
                if birth_month < 1 or birth_month > 12:
                    return False, f"出生月份 {birth_month} 不合理"
                if birth_day < 1 or birth_day > 31:
                    return False, f"出生日期 {birth_day} 不合理"
                    
                # 检查2月份日期
                if birth_month == 2:
                    is_leap = (birth_year % 4 == 0 and birth_year % 100 != 0) or (birth_year % 400 == 0)
                    if (is_leap and birth_day > 29) or (not is_leap and birth_day > 28):
                        return False, f"{birth_year}年2月没有{birth_day}日"
                        
                # 检查小月
                if birth_month in [4, 6, 9, 11] and birth_day > 30:
                    return False, f"{birth_month}月没有{birth_day}日"
            except:
                return False, "出生日期格式不正确"
                
            return True, ""
        else:
            return False, "身份证号码长度应为15位或18位"

    def convert_markdown_to_plain(self, markdown_text):
        """将Markdown文本转换为易读的纯文本格式"""
        # 替换标题
        lines = markdown_text.split('\n')
        plain_text = []
        
        for line in lines:
            # 处理标题行
            if line.startswith('# '):  # 一级标题
                plain_text.append("\n" + line[2:] + "\n" + "="*len(line[2:]))
            elif line.startswith('## '):  # 二级标题
                plain_text.append("\n" + line[3:] + "\n" + "-"*len(line[3:]))
            elif line.startswith('### '):  # 三级标题
                plain_text.append("\n" + line[4:])
            # 处理列表项
            elif line.strip().startswith('- '):
                plain_text.append("• " + line.strip()[2:])
            elif line.strip().startswith('* '):
                plain_text.append("• " + line.strip()[2:])
            # 处理有序列表
            elif re.match(r'^\d+\.\s', line.strip()):
                plain_text.append(line.strip())
            # 处理粗体文本
            elif '**' in line:
                # 将**text**替换为text
                processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                plain_text.append(processed_line)
            else:
                plain_text.append(line)
        
        # 合并处理后的行
        return '\n'.join(plain_text)
    
    def show_readme(self):
        """显示README文件内容"""
        readme_path = os.path.join(self.app_dir, "README.md")
        
        # 如果README文件不存在，创建默认内容
        if not os.path.exists(readme_path):
            with open(readme_path, "w", encoding="utf-8") as f:
                f.write("""# 情况说明生成器 V1.11

## 软件介绍
情况说明生成器是一款用于生成集体土地安置房分配协议的专业工具，专为房产安置业务设计。

## 主要功能
1. 家庭成员管理：添加、删除家庭成员及其身份证信息，第一位成员自动标记为户主
2. 身份证验证：自动校验身份证格式和有效性，避免录入错误
3. 项目信息录入：包含拆迁项目名称和安置小区名称
4. 安置房信息管理：支持多套房产信息的录入与管理
5. 产权人设置：每套房产可设置多个产权人
6. 文档生成：自动生成标准格式的Word文档
7. 直接打印：支持无需打开Word直接打印文档

## 操作技巧
1. 使用回车键在表单中快速导航，类似Excel的操作方式
2. 添加多个产权人时，可点击"+"按钮为同一套房产添加多个所有者
3. 身份证输入完成后，系统会自动验证，错误会有具体提示""")
        
        # 读取README文件内容
        try:
            with open(readme_path, "r", encoding="utf-8") as f:
                readme_content = f.read()
            
            # 将Markdown转换为易读的纯文本
            plain_content = self.convert_markdown_to_plain(readme_content)
            
            # 创建自定义对话框
            readme_dialog = tk.Toplevel(self.root)
            readme_dialog.title("情况说明生成器 - 使用说明")
            readme_dialog.geometry("700x500")
            readme_dialog.transient(self.root)
            readme_dialog.grab_set()
            
            # 创建文本区域
            text_frame = ttk.Frame(readme_dialog, padding=10)
            text_frame.pack(fill=tk.BOTH, expand=True)
            
            # 添加滚动条
            scrollbar = ttk.Scrollbar(text_frame)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            # 创建文本控件
            text = tk.Text(text_frame, wrap=tk.WORD, yscrollcommand=scrollbar.set, font=("黑体", 12))
            text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.config(command=text.yview)
            
            # 插入内容
            text.insert(tk.END, plain_content)
            text.config(state=tk.DISABLED)  # 设为只读
            
            # 添加关闭按钮
            btn_frame = ttk.Frame(readme_dialog, padding=10)
            btn_frame.pack(fill=tk.X)
            
            close_btn = ttk.Button(btn_frame, text="关闭", command=readme_dialog.destroy, width=15)
            close_btn.pack(side=tk.RIGHT)
            
        except Exception as e:
            messagebox.showerror("错误", f"无法读取README文件: {str(e)}")

if __name__ == "__main__":
    try:
        # 如果是Windows系统，尝试隐藏控制台窗口
        if sys.platform == 'win32':
            import ctypes
            ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
    except:
        pass
        
    root = tk.Tk()
    app = SituationReportGenerator(root)
    root.mainloop()